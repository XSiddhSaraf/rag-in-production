"""Document processing for PDF and Word files."""

from typing import List, Tuple
from pathlib import Path
import PyPDF2
import pdfplumber
from docx import Document
from loguru import logger
import re


class DocumentProcessor:
    """Process PDF and Word documents."""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """Initialize document processor.
        
        Args:
            chunk_size: Target size for text chunks
            chunk_overlap: Overlap between consecutive chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        try:
            # Try pdfplumber first (better for complex layouts)
            with pdfplumber.open(file_path) as pdf:
                text_parts = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                
                if text_parts:
                    return "\n\n".join(text_parts)
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {e}. Trying PyPDF2...")
        
        # Fallback to PyPDF2
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_parts = []
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                
                return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            raise ValueError(f"Failed to extract text from PDF: {e}")
    
    def extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from Word document.
        
        Args:
            file_path: Path to Word file
            
        Returns:
            Extracted text
        """
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        text_parts.append(row_text)
            
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"Word document extraction failed: {e}")
            raise ValueError(f"Failed to extract text from Word document: {e}")
    
    def extract_text(self, file_path: Path) -> str:
        """Extract text from document (auto-detect format).
        
        Args:
            file_path: Path to document
            
        Returns:
            Extracted text
        """
        suffix = file_path.suffix.lower()
        
        if suffix == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif suffix in ['.docx', '.doc']:
            return self.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")
    
    def clean_text(self, text: str) -> str:
        """Clean extracted text.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters that may interfere
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\xff]', '', text)
        # Normalize line breaks
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        return text.strip()
    
    def chunk_text(self, text: str) -> List[str]:
        """Split text into overlapping chunks.
        
        Args:
            text: Text to chunk
            
        Returns:
            List of text chunks
        """
        # Split into sentences (basic approach)
        sentences = re.split(r'(?<=[.!?])\s+', text)
        
        chunks = []
        current_chunk = []
        current_size = 0
        
        for sentence in sentences:
            sentence_size = len(sentence)
            
            if current_size + sentence_size > self.chunk_size and current_chunk:
                # Save current chunk
                chunks.append(' '.join(current_chunk))
                
                # Start new chunk with overlap
                overlap_text = ' '.join(current_chunk)
                if len(overlap_text) > self.chunk_overlap:
                    # Keep last portion for overlap
                    overlap_sentences = []
                    overlap_size = 0
                    for sent in reversed(current_chunk):
                        if overlap_size + len(sent) <= self.chunk_overlap:
                            overlap_sentences.insert(0, sent)
                            overlap_size += len(sent)
                        else:
                            break
                    current_chunk = overlap_sentences
                    current_size = overlap_size
                else:
                    current_chunk = []
                    current_size = 0
            
            current_chunk.append(sentence)
            current_size += sentence_size
        
        # Add final chunk
        if current_chunk:
            chunks.append(' '.join(current_chunk))
        
        return chunks
    
    def process_document(self, file_path: Path) -> Tuple[str, List[str]]:
        """Process document: extract, clean, and chunk.
        
        Args:
            file_path: Path to document
            
        Returns:
            Tuple of (full_text, chunks)
        """
        logger.info(f"Processing document: {file_path}")
        
        # Extract and clean
        raw_text = self.extract_text(file_path)
        clean_text = self.clean_text(raw_text)
        
        # Chunk
        chunks = self.chunk_text(clean_text)
        
        logger.info(f"Extracted {len(clean_text)} characters, created {len(chunks)} chunks")
        
        return clean_text, chunks


# Global document processor instance
document_processor = DocumentProcessor()
